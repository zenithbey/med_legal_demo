import fitz
from pdf2image import convert_from_bytes
import pytesseract
from docx import Document
import io
from openai import OpenAI
import os
from pathlib import Path
from dotenv import load_dotenv
from datetime import datetime
import magic
import re

env_path = Path(__file__).resolve().parent.parent / "config.env"
load_dotenv(env_path)

client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

REQUIRED_DOC_TYPES = [
    "incident_report", 
    "medical_history", 
    "treatment_plan",
    "discharge_summary"
]

def extract_pdf_text(file_stream):
    """Extract text from searchable PDFs"""
    text = ""
    try:
        with fitz.open(stream=file_stream.read(), filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
        return text if text.strip() else None
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {str(e)}")

def is_scanned_pdf(content):
    """Check if PDF needs OCR"""
    try:
        with fitz.open(stream=content, filetype="pdf") as doc:
            return not any(page.get_text().strip() for page in doc)
    except Exception as e:
        raise ValueError(f"PDF scan check failed: {str(e)}")

def ocr_pdf(file_stream):
    """Process scanned PDFs with OCR"""
    images = convert_from_bytes(file_stream.read())
    return "\n".join(pytesseract.image_to_string(img) for img in images)

def extract_docx_text(file_stream):
    """Extract text from DOCX files"""
    try:
        doc = Document(io.BytesIO(file_stream.read()))
        full_text = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            full_text.extend(' | '.join(cell.text for cell in row.cells) 
                          for row in table.rows)
        return '\n'.join(filter(None, full_text))
    except Exception as e:
        raise ValueError(f"DOCX processing failed: {str(e)}")

def process_document(file, true_mime):
    """Main document processing function"""
    content = file.read()
    file.stream.seek(0)
    filename = (file.filename or "unnamed").lower()

    try:
        # PDF handling
        if true_mime == 'application/pdf':
            if is_scanned_pdf(content):
                text = ocr_pdf(io.BytesIO(content))
            else:
                text = extract_pdf_text(io.BytesIO(content))
        
        # DOCX handling
        elif true_mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            text = extract_docx_text(io.BytesIO(content))
        
        else:
            raise ValueError(f"Unsupported file type: {true_mime}")

        if not text.strip():
            raise ValueError("Document appears empty or unreadable")

        return {
            "filename": filename,
            "content": text[:1000] + "...",
            "analysis": analyze_with_gpt4(text) if text else "Analysis failed: Empty document",
            "mime_type": true_mime,
            "processed_at": datetime.now().isoformat(),
            "status": "success"
        }

    except Exception as e:
        return {
            "filename": filename,
            "analysis": f"Analysis failed: {str(e)}",
            "status": "failed"
        }


def parse_chronology(analysis_text):
    """Convert AI output to structured timeline"""
    events = []
    date_pattern = r"\d{4}-\d{2}-\d{2}"
    
    for line in analysis_text.split('\n'):
        if "Date:" in line and "Event:" in line:
            date_match = re.search(date_pattern, line)
            if date_match:
                event = {
                    'date': datetime.strptime(date_match.group(), "%Y-%m-%d").date(),
                    'description': line.split("Event:")[-1].strip()
                }
                events.append(event)
    return sorted(events, key=lambda x: x['date'])

def analyze_with_gpt4(text):
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {
                "role": "system",
                "content": """Analyze medical documents and extract:
                1. Chronological events (date in YYYY-MM-DD format, description)
                2. Document type classification
                3. Missing required document types
                4. Potential liability factors. Use markdown formatting."""
            },
            {
                "role": "user",
                "content": text[:15000]  # Limit context size
            }
        ]
    )
    return response.choices[0].message.content

def detect_missing_reports(analyses):
    """Identify missing document types from analysis texts"""
    required_docs = [
        "Accident/Incident Report",
        "Hospital Admission Records",
        "Discharge Summary",
        "Imaging Results",
        "Treatment Plans",
        "Insurance Claim Forms"
    ]
    
    # Extract analysis texts whether we get objects or strings
    analysis_texts = []
    for item in analyses:
        if isinstance(item, dict):
            analysis_texts.append(item.get('text', item.get('analysis', '')))
        else:
            analysis_texts.append(str(item))
    
    analysis_text = "\n\n".join(analysis_texts)
    
    prompt = f"""Analyze these medical documents for missing required documents:
    
    Required Documents:
    {chr(10).join(required_docs)}
    
    Document Content:
    {analysis_text[:12000]}
    
    Return ONLY a comma-separated list of missing document types from the required list.
    Example: 'Accident/Incident Report, Imaging Results'
    """
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "Identify missing medical documents from analysis"},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1
        )
        
        missing = [x.strip() for x in response.choices[0].message.content.split(",")]
        return [doc for doc in required_docs if any(
            m.lower() in doc.lower() for m in missing
        )]
    
    except Exception as e:
        print(f"Error in detect_missing_reports: {str(e)}")
        return []

# Malpractice Reviewer

# Specialty-specific resources map
SPECIALTY_RESOURCES = {
    "Obstetrics": (
        "Williams Obstetrics, Gabbe's Obstetrics, "
        "Creasy and Resnik's Maternal-Fetal Medicine, "
        "The Johns Hopkins Manual of Gynecology and Obstetrics"
    ),
    "Emergency Medicine": (
        "Rosen's Emergency Medicine, Tintinalli's Emergency Medicine"
    ),
    "Stroke/Neurology": (
        "Stroke: Pathophysiology, Diagnosis, and Management (Mohr), "
        "Stroke Medicine (Oxford Specialist Handbooks in Neurology)"
    ),
    "Oncology": (
        "The ESMO Handbook of Cancer Diagnosis and Treatment Evaluation (2nd Ed), "
        "The American Cancer Society's Principles of Oncology, "
        "The Cancer Journey by Dr. Chadi Nabhan"
    ),
    "ENT": (
        "Advanced Health Assessment and Diagnostic Reasoning (3rd ed), "
        "Bailey's Head and Neck Surgery: Otolaryngology (5th ed), "
        "Atlas of Endoscopic Sinus and Skull Base Surgery"
    ),
    "General Surgery": "Schwartz's Principles of Surgery, ACS Surgery Principles and Practice"
}

def extract_document_text(file, true_mime):
    """Extract full text without analysis"""
    content = file.read()
    file.stream.seek(0)
    filename = (file.filename or "unnamed").lower()

    # PDF handling
    if true_mime == 'application/pdf':
        if is_scanned_pdf(content):
            text = ocr_pdf(io.BytesIO(content))
        else:
            text = extract_pdf_text(io.BytesIO(content))
    
    # DOCX handling
    elif true_mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        text = extract_docx_text(io.BytesIO(content))
    
    else:
        raise ValueError(f"Unsupported file type: {true_mime}")

    if not text.strip():
        raise ValueError("Document appears empty or unreadable")
        
    return text

def analyze_standard_of_care(specialty, text):
    """Analyze medical records for SOC compliance and causation"""
    resources = SPECIALTY_RESOURCES.get(specialty, "Standard medical textbooks and guidelines")
    
    prompt = f"""
    You are an expert physician analyzing a medical case to evaluate:
    1. Adherence to the standard of care (SOC)
    2. Causation of injuries based on breaches of SOC
    3. Potential liability factors
    
    Apply a systematic, evidence-based approach using differential diagnosis methodology.
    
    Specialty: {specialty}
    Core Resources: {resources}
    
    Case Documentation:
    {text[:15000]}
    
    Provide a comprehensive analysis with the following sections:
    
    ### Case Summary
    - Concise overview of key events and timeline
    
    ### Standard of Care Evaluation
    - SOC benchmarks applicable to this case
    - Specific breaches of SOC (if any)
    - Adherence to SOC where appropriate
    - Use bullet points with ✅ for adherence and ❌ for breaches
    
    ### Causation Analysis
    - Correlation between identified breaches and patient outcomes
    - Alternative explanations for outcomes (differential diagnosis)
    - Strength of causal relationship
    
    ### Liability Assessment
    - Key liability factors
    - Potential damages calculation considerations
    - Contributing factors (systemic, communication, etc.)
    
    ### Supporting References
    - Relevant clinical literature and guidelines
    - Specific references to core resources
    - PubMed/Google Scholar references where applicable
    
    Format your response using Markdown. Be thorough but concise.
    """
    
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": "You are a medical legal expert analyzing standard of care compliance."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.2,
        max_tokens=3000
    )
    
    return response.choices[0].message.content

__all__ = ['process_document', 'detect_missing_reports'] 
